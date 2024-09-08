import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

# Step 1: Load the Excel file
# Replace 'your_excel_file.xlsx' with the path to your actual Excel file
excel_file_path = 'your_excel_file.xlsx'
# Replace 'your_sheet_name' with the name of the sheet containing the data
df = pd.read_excel(excel_file_path, sheet_name='your_sheet_name', skiprows=5)

# Step 2: Create the output folder for PDFs (if it doesn't exist)
output_folder = 'generated_pdfs'
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# Step 3: Loop through each row of data
# Update the range to match the rows you want to process
for idx in range(1, len(df)):
    data_row = df.iloc[idx]

    # Step 4: Extract and clean the data for this row
    # Remove leading/trailing spaces
    # Adjust column name if necessary
    full_location = str(data_row['Unnamed: 3']).strip()
    # Extract the first part of the location (e.g., "Sidoarjo")
    place_name = full_location.split(",")[0]
    # Adjust column name if necessary
    date_string = str(data_row['Unnamed: 7'])
    formatted_date = pd.to_datetime(
        date_string).strftime("%d %B %Y")  # Format date
    # Combine place and date
    place_and_date = f"{place_name}, {formatted_date}"

    # Step 5: Define replacements based on the current row
    replacements = {
        "Name": str(data_row['Unnamed: 1']),  # Adjust column name if necessary
        # Adjust column name if necessary
        "Street address": str(data_row['Unnamed: 2']),
        "Postcode location": full_location,
        # Adjust column name if necessary
        "Country": str(data_row['Unnamed: 4']),
        # Adjust column name if necessary
        "Phone number": str(data_row['Unnamed: 5']),
        # Adjust column name if necessary
        "Recipient of the UCO (Collecting Point)": str(data_row['Unnamed: 6']),
        "Place date": place_and_date
    }

    # Step 6: Load the Word template
    # Replace 'your_word_template.docx' with the path to your Word template file
    word_template_path = 'your_word_template.docx'
    doc = Document(word_template_path)

    # Step 7: Replace data in specific table cells and adjust font size/alignment
    table = doc.tables[0]
    table.cell(1, 1).text = replacements["Name"]
    table.cell(2, 1).text = replacements["Street address"]
    table.cell(3, 1).text = replacements["Postcode location"]
    table.cell(4, 1).text = replacements["Country"]
    table.cell(5, 1).text = replacements["Phone number"]

    # Adjust font size and alignment for "Recipient of the UCO"
    recipient_cell = table.cell(8, 1)
    recipient_cell.text = replacements["Recipient of the UCO (Collecting Point)"]
    recipient_cell_paragraph = recipient_cell.paragraphs[0]
    recipient_cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in recipient_cell_paragraph.runs:
        run.font.size = Pt(16)  # Set the font size to 16 points

    # Adjust font size and alignment for "Place, date"
    place_date_cell = table.cell(11, 0)
    place_date_cell.text = replacements["Place date"]
    place_date_cell_paragraph = place_date_cell.paragraphs[0]
    place_date_cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in place_date_cell_paragraph.runs:
        run.font.size = Pt(19.5)  # Set the font size to 19.5 points

    # Step 8: Generate the PDF file name with format "1. Name.pdf"
    name_part = replacements["Name"].replace(
        "/", "-").replace("\\", "-")  # Clean file name
    pdf_name = f"{idx}. {name_part}.pdf"
    output_pdf_path = os.path.join(output_folder, pdf_name)

    # Step 9: Convert the Word document directly to PDF without saving DOCX
    doc.save('temp_doc.docx')  # Temporary save to convert to PDF
    convert('temp_doc.docx', output_pdf_path)  # Convert to PDF
    os.remove('temp_doc.docx')  # Remove the temporary DOCX file

    print(f"Generated PDF for row {idx} with name '{pdf_name}'")

print("Automation complete. All PDFs have been generated.")
